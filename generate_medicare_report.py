import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize Document
doc = Document()

# Page setup (margins)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("Medicare AI: Personalized Healthcare Recommendation System")
run_title.font.name = 'Arial'
run_title.font.size = Pt(20)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(18, 53, 91)

# Subtitle / Metadata
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run(
    "Zidio Development Internship Project\n"
    "Developed By: GULAFSHA\n"
    "Project Duration: August 2026 – September 2026 (1 Month)\n"
)
run_sub.font.name = 'Arial'
run_sub.font.size = Pt(11.5)
run_sub.font.color.rgb = RGBColor(100, 116, 139)

# Report Sections
sections_data = [
    ("1. Project Overview & Vision", 
     "Medicare AI is a personalized healthcare recommendation and disease diagnosis platform developed as a Zidio Development internship project. The platform is designed to bridge the gap between advanced machine learning predictive models and accessible personal health management. By evaluating patient demographics, vitals (blood pressure, cholesterol, age, gender), and symptom profiles, the system provides accurate disease predictions, dataset-driven risk stratification, and targeted pharmaceutical recommendations."),
    
    ("2. Vision & Objectives",
     "• Accurate Disease Prediction: Multi-symptom pattern recognition mapped against 116 distinct medical conditions using trained machine learning models.\n"
     "• Automated Risk Stratification: Categorizing patient health risk levels into Low, Medium, and High based on clinical parameters and dataset patterns.\n"
     "• Personalized Treatment & Care: Providing precise pharmaceutical drug names, exact dosages, dietary guidelines, and lifestyle precautions.\n"
     "• Secure User Authentication & Dashboard: Implementing Flask-Login and SQLAlchemy to maintain user sessions, login metrics, and historical medical diagnosis reports.\n"
     "• Developer Telemetry & Admin Panel: Offering a dedicated admin command center to monitor global site visits, user engagement, and system health checks."),
    
    ("3. Technical Methodology & System Architecture",
     "The project follows a rigorous end-to-end data science and full-stack web development workflow:\n"
     "• Phase 1: Data Ingestion & ETL Pipeline — Cleaning and formatting Cleaned_Dataset.csv, handling missing values, and mapping multi-symptom feature vectors.\n"
     "• Phase 2: Feature Engineering & Vitals Weighting — Integrating patient demographics and clinical vitals (blood pressure, cholesterol level, age, gender) into the predictive feature array.\n"
     "• Phase 3: Model Development & Validation — Training supervised classification models (evaluated via scikit-learn metrics) combined with probability scaling for confidence calculation.\n"
     "• Phase 4: Full-Stack Integration & Deployment — Developing a responsive Flask web framework paired with SQLite database storage (medicare.db), modern CSS styling, and interactive UI components."),
    
    ("4. AI Techniques & Algorithms",
     "• Supervised Classification: Machine learning classification models combined with label encoders (disease_encoder.pkl) to output top-5 predicted conditions.\n"
     "• Risk Level Assessment: Dynamic mapping utilizing dataset risk classifications and clinical overrides (elevated blood pressure, high cholesterol, or advanced age automatically escalate risk levels to High).\n"
     "• Relational Treatment Lookup: Structured medical database mapping conditions to verified pharmaceutical names and dosages."),
    
    ("5. Technology Stack",
     "• Programming Language: Python 3.12\n"
     "• Backend Framework: Flask, Flask-SQLAlchemy, Flask-Login, Flask-CORS\n"
     "• Machine Learning & Data Science: Scikit-learn, Pandas, NumPy, Joblib\n"
     "• Database: SQLite (medicare.db)\n"
     "• Frontend: HTML5, CSS3 (Inter & Poppins typography, responsive grid layouts, custom UI risk badges)\n"
     "• Version Control & Deployment: Git & GitHub"),
    
    ("6. Conclusion & Future Roadmap",
     "Medicare AI successfully demonstrates the practical implementation of machine learning in healthcare technology. Developed during the August 2026 – September 2026 internship timeline under Zidio Development, the platform delivers a robust, secure, and data-driven diagnostic tool. Future roadmap goals include real-time EHR integration, advanced recommendation systems, and cloud-native scaling.")
]

for heading, text in sections_data:
    h = doc.add_paragraph()
    run_h = h.add_run(heading)
    run_h.font.name = 'Arial'
    run_h.font.size = Pt(13)
    run_h.font.bold = True
    run_h.font.color.rgb = RGBColor(22, 119, 255)
    
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)
    run_p = p.runs[0]
    run_p.font.name = 'Arial'
    run_p.font.size = Pt(10.5)
    run_p.font.color.rgb = RGBColor(31, 41, 55)

# Save document
output_filename = "Medicare_AI_Zidio_Project_Report.docx"
doc.save(output_filename)
print(f"SUCCESS: {output_filename} generated successfully in your project folder!")