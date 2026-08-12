import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize Document
doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("MediCare AI — Comprehensive Project Report")
run_title.font.name = 'Arial'
run_title.font.size = Pt(22)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(18, 53, 91)

# Subtitle
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("Production-Grade Healthcare Analytics & Machine Learning Platform\n")
run_sub.font.name = 'Arial'
run_sub.font.size = Pt(13)
run_sub.font.color.rgb = RGBColor(100, 116, 139)

# Report Sections
sections_data = [
    ("1. Executive Summary", 
     "MediCare AI is an advanced, production-grade artificial intelligence and machine learning healthcare analytics platform designed to empower clinicians and patients with accurate disease predictions, precise risk stratification, and targeted pharmaceutical recommendations. Developed for clinical and personal health evaluation, the platform integrates machine learning classification models with clinical parameter analysis (such as age, gender, blood pressure, and cholesterol levels) to assess medical conditions. By leveraging structured historical medical datasets (Cleaned_Dataset.csv), MediCare AI reduces diagnostic ambiguity and ensures reliable, data-driven treatment pathways."),
    
    ("2. Project Objectives",
     "The primary goal of this project was to build a robust, secure, full-stack data science web application addressing core clinical challenges:\n"
     "• Accurate Disease Prediction: Multi-symptom pattern recognition mapped against over 100 distinct medical conditions.\n"
     "• Risk Stratification: Automated classification of patient health risk levels (Low, Medium, High) incorporating vital signs and demographic metrics.\n"
     "• Personalized Treatment & Medication Database: Providing exact pharmaceutical drug names, dosages, dietary guidelines, and precautions.\n"
     "• User Health Tracking: Secure user authentication (Flask-Login), session tracking, and an interactive customer dashboard storing historical medical reports.\n"
     "• Developer Admin Telemetry: A dedicated admin command center tracking total platform visits, registered user accounts, login frequencies, and system performance metrics."),
    
    ("3. Technical Methodology",
     "The project was executed in a structured four-phase approach:\n"
     "• Phase 1: Data Ingestion & ETL Pipeline — Ingested and cleaned the clinical health dataset (Cleaned_Dataset.csv), handling missing values, standardizing categorical variables (blood pressure, cholesterol, gender), and mapping symptom feature vectors.\n"
     "• Phase 2: Feature Engineering & Vitals Weighting — Engineered multi-symptom input arrays combined with patient vitals (age, gender, blood pressure, cholesterol_level) to ensure physical parameters actively influence disease prediction and risk calculation.\n"
     "• Phase 3: Model Development & Validation — Trained supervised classification models evaluated using industry-standard metrics (accuracy, precision, recall, and prediction probabilities).\n"
     "• Phase 4: Full-Stack Integration & Deployment — Integrated the backend logic into a responsive Flask web framework paired with SQLite database storage (Flask-SQLAlchemy), modern CSS UI design, and interactive user dashboards."),
    
    ("4. AI Techniques & Algorithms",
     "4.1 Disease Classification & Multi-Symptom Matching\n"
     "• Technique: Supervised Machine Learning Classification (best_model.pkl) with Label Encoding (disease_encoder.pkl).\n"
     "• Mechanism: The model analyzes symptom feature vectors (feature_columns.pkl) alongside patient vitals to output top predicted conditions with computed confidence percentages.\n\n"
     "4.2 Risk Stratification Analysis (Low, Medium, High)\n"
     "• Technique: Rule-Based & Dataset-Driven Clinical Scoring.\n"
     "• Mechanism: Integrates direct risk classifications from the cleaned medical dataset (covering 116 unique conditions across Low, Medium, and High risk profiles) with clinical overrides. Elevated blood pressure (blood_pressure == 2), high cholesterol (cholesterol_level == 2), or advanced age (age > 60) automatically escalate patient risk assessments to HIGH for clinical safety.\n\n"
     "4.3 Precise Medication & Treatment Mapping\n"
     "• Technique: Relational Healthcare Database Lookup.\n"
     "• Mechanism: Maps predicted conditions directly to structured treatment dictionaries containing exact pharmaceutical drug names, dosages, dietary restrictions, and workout precautions."),
    
    ("5. Technology Stack",
     "• Programming Language: Python 3.12\n"
     "• Web Framework: Flask, Flask-SQLAlchemy, Flask-Login, Flask-CORS\n"
     "• Machine Learning & Data Science: Scikit-learn, Pandas, NumPy, Joblib\n"
     "• Database: SQLite (medicare.db)\n"
     "• Frontend: HTML5, CSS3 (Inter & Poppins typography, responsive grid layouts, custom card components)\n"
     "• Visualization & Telemetry: Chart.js (Admin analytics and graphs)"),
    
    ("6. Conclusion",
     "MediCare AI successfully demonstrates the practical application of machine learning in healthcare technology. By combining rigorous classification models with clinical parameter weighting and a secure user dashboard, the platform provides a scalable, professional solution that delivers measurable value for personal health tracking and clinical decision support.")
]

for heading, text in sections_data:
    h = doc.add_paragraph()
    run_h = h.add_run(heading)
    run_h.font.name = 'Arial'
    run_h.font.size = Pt(14)
    run_h.font.bold = True
    run_h.font.color.rgb = RGBColor(22, 119, 255)
    
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(10)
    run_p = p.runs[0]
    run_p.font.name = 'Arial'
    run_p.font.size = Pt(11)
    run_p.font.color.rgb = RGBColor(31, 41, 55)

# Save document
doc.save("MediCare_AI_Project_Report.docx")
print("SUCCESS: MediCare_AI_Project_Report.docx created successfully in your project folder!")